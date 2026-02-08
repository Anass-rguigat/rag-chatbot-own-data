"""
RAG Chatbot with Memory.

A retrieval-augmented generation (RAG) chatbot that answers only from user-provided
documents (PDF, TXT, MD, DOCX) in a single knowledge-base folder. Supports two modes:
FAST (laptop, low resource) and COMPLETE (server/API, full quality). Uses Streamlit
for the UI, LangChain for the pipeline, FAISS for the vector store, and Groq or
Ollama for the LLM.
"""

from typing import List, Optional, Tuple
import os
import json
import hashlib
import time
import warnings
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeoutError

from langchain_core.messages import HumanMessage

warnings.filterwarnings("ignore", message=".*migrating_memory.*")
try:
    from langchain_core._api.deprecation import LangChainDeprecationWarning
    warnings.filterwarnings("ignore", category=LangChainDeprecationWarning)
except Exception:
    pass

from dotenv import load_dotenv
load_dotenv()

import streamlit as st
from langchain_groq import ChatGroq
from langchain_community.chat_models.ollama import ChatOllama
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.chains import ConversationalRetrievalChain
from langchain_community.chat_message_histories import ChatMessageHistory
from langchain.memory import ConversationBufferMemory
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from pypdf import PdfReader
from langchain.prompts import PromptTemplate

SUPPORTED_EXTENSIONS = (".pdf", ".txt", ".md", ".docx")

# COMPLETE mode prompt (Server/API) — long, strict, 100% correct and complete
QA_PROMPT_TEMPLATE = """You are the user's reliable source. Your answer must be 100% CORRECT (only from Context) and 100% COMPLETE (every relevant detail from Context).

STRICT RULES:
1. Use ONLY information that appears in the Context below. Do not add, infer, or invent anything. Zero hallucination.
2. Give a COMPLETE answer: include EVERY relevant detail, name, number, date, and fact from the Context that answers the question. Do not summarize away or skip important information.
3. For lists (e.g. "list all...", "what are the...", "quels sont..."): list EVERY relevant item in the Context. Each item once; use one clear form if the same thing appears under different names. Do not omit any item.
4. For overview / "what is this about?": state the document type and topic, then mention EVERY distinct section, title, or main item that appears in the Context.
5. Prefer exact wording and figures from the Context so the answer is accurate.
6. If the answer is not in the Context, say exactly: "This does not appear in the excerpts I was given."
7. Before finishing: verify you did not omit ANY relevant part of the Context that answers the question. If in doubt, include it.

Context:
{context}

Question: {question}

Answer (100% correct and 100% complete, from the Context only):"""
QA_PROMPT = PromptTemplate(
    template=QA_PROMPT_TEMPLATE,
    input_variables=["context", "question"],
)

# FAST mode prompt (Laptop) — short, strict, still no hallucination
QA_PROMPT_FAST_TEMPLATE = """Answer ONLY from the Context below. Keep it short and precise.
If the answer is not in the Context, say exactly: "This does not appear in the excerpts I was given."

Context:
{context}

Question: {question}

Answer:"""
QA_PROMPT_FAST = PromptTemplate(
    template=QA_PROMPT_FAST_TEMPLATE,
    input_variables=["context", "question"],
)

# ============================================================================
# Configuration (all values from .env; defaults only when env var is unset)
# ============================================================================

def _env(key: str, default: str) -> str:
    return os.getenv(key, default).strip()


def _env_int(key: str, default: str) -> int:
    return int(_env(key, default))


def _env_float(key: str, default: str) -> float:
    return float(_env(key, default))


DATA_FOLDER_PATH = _env("DATA_FOLDER_PATH", "") or _env("PDF_FOLDER_PATH", "./fin_ed_docs")
GROQ_API_KEY = _env("GROQ_API_KEY", "your-groq-api-key-here")
OLLAMA_BASE_URL = _env("OLLAMA_BASE_URL", "http://localhost:11434")
EMBEDDING_MODEL = _env("EMBEDDING_MODEL", "sentence-transformers/all-MiniLM-L6-v2")
MAX_SOURCE_PREVIEW = _env_int("MAX_SOURCE_PREVIEW", "800")
TEMPERATURE = _env_float("TEMPERATURE", "0")

# COMPLETE mode (Server/API): used when Mode = COMPLETE (Server/API)
CHUNK_SIZE_COMPLETE = _env_int("CHUNK_SIZE", "2000")
CHUNK_OVERLAP_COMPLETE = _env_int("CHUNK_OVERLAP", "450")
RETRIEVER_K_COMPLETE = _env_int("RETRIEVER_K", "15")
RETRIEVER_FETCH_K_COMPLETE = _env_int("RETRIEVER_FETCH_K", "40")
GROQ_MODEL_COMPLETE = _env("GROQ_MODEL", "llama-3.3-70b-versatile")
GROQ_MAX_TOKENS_COMPLETE = _env_int("GROQ_MAX_TOKENS", "8192")
OLLAMA_MODEL_COMPLETE = _env("OLLAMA_MODEL", "tinyllama")
OLLAMA_NUM_PREDICT_COMPLETE = _env_int("OLLAMA_NUM_PREDICT", "2048")
LLM_TIMEOUT_COMPLETE = _env_int("LLM_TIMEOUT_SEC", "90")

# FAST mode (Laptop): used when Mode = FAST (Laptop); separate env vars
CHUNK_SIZE_FAST = _env_int("CHUNK_SIZE_FAST", "800")
CHUNK_OVERLAP_FAST = _env_int("CHUNK_OVERLAP_FAST", "120")
RETRIEVER_K_FAST = _env_int("RETRIEVER_K_FAST", "5")
RETRIEVER_FETCH_K_FAST = _env_int("RETRIEVER_FETCH_K_FAST", "15")
GROQ_MODEL_FAST = _env("GROQ_MODEL_FAST", "llama-3.1-8b-instant")
GROQ_MAX_TOKENS_FAST = _env_int("GROQ_MAX_TOKENS_FAST", "512")
OLLAMA_MODEL_FAST = _env("OLLAMA_MODEL_FAST", "phi3:mini")
OLLAMA_NUM_PREDICT_FAST = _env_int("OLLAMA_NUM_PREDICT_FAST", "256")
LLM_TIMEOUT_FAST = _env_int("LLM_TIMEOUT_FAST", "60")

GROQ_RETRY_DELAYS = [2, 5, 10]

VECTORSTORE_CACHE_PATH = _env("VECTORSTORE_CACHE_PATH", "./.vectorstore")
VECTORSTORE_CACHE_PATH_FAST = _env("VECTORSTORE_CACHE_PATH_FAST", "./.vectorstore_fast")
EMBEDDING_DEVICE = _env("EMBEDDING_DEVICE", "auto")


def _data_folder_signature(data_folder_path: str) -> str:
    """Return a hash that changes when any supported document in the folder is added/removed/changed."""
    if not os.path.isdir(data_folder_path):
        return ""
    parts = []
    for f in sorted(os.listdir(data_folder_path)):
        if not f.lower().endswith(SUPPORTED_EXTENSIONS):
            continue
        path = os.path.join(data_folder_path, f)
        try:
            stat = os.stat(path)
            parts.append(f"{f}|{stat.st_size}|{stat.st_mtime}")
        except OSError:
            parts.append(f"{f}|")
    return hashlib.sha256(json.dumps(parts, sort_keys=True).encode()).hexdigest()


def _vectorstore_manifest_path(cache_path: str) -> str:
    return os.path.join(cache_path, "manifest.json")


def _try_load_cached_vectorstore(embeddings, cache_path: str, chunk_size: int, chunk_overlap: int) -> Optional[FAISS]:
    """Load FAISS from disk if cache exists and matches current documents and settings."""
    manifest_path = _vectorstore_manifest_path(cache_path)
    if not os.path.isfile(manifest_path):
        return None
    try:
        with open(manifest_path, "r", encoding="utf-8") as f:
            manifest = json.load(f)
    except (json.JSONDecodeError, OSError):
        return None
    current_sig = _data_folder_signature(DATA_FOLDER_PATH)
    cached_sig = manifest.get("data_signature") or manifest.get("pdf_signature")
    if cached_sig != current_sig:
        return None
    if manifest.get("embedding_model") != EMBEDDING_MODEL:
        return None
    if manifest.get("chunk_size") != chunk_size or manifest.get("chunk_overlap") != chunk_overlap:
        return None
    index_path = os.path.join(cache_path, "index.faiss")
    if not os.path.isfile(index_path):
        return None
    try:
        return FAISS.load_local(
            cache_path,
            embeddings,
            allow_dangerous_deserialization=True,
        )
    except Exception:
        return None


def _save_vectorstore_cache(vectorstore: FAISS, cache_path: str, chunk_size: int, chunk_overlap: int) -> None:
    """Save FAISS and manifest so next run can load from disk."""
    os.makedirs(cache_path, exist_ok=True)
    vectorstore.save_local(cache_path)
    manifest = {
        "data_signature": _data_folder_signature(DATA_FOLDER_PATH),
        "embedding_model": EMBEDDING_MODEL,
        "chunk_size": chunk_size,
        "chunk_overlap": chunk_overlap,
    }
    with open(_vectorstore_manifest_path(cache_path), "w", encoding="utf-8") as f:
        json.dump(manifest, f, indent=2)


def _load_one_document(file_path: str) -> List[Document]:
    """Load a single file (PDF, TXT, MD, or DOCX) into LangChain Documents."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        reader = None
        try:
            reader = PdfReader(file_path)
            pdf_page_count = len(reader.pages)
            loader = PyPDFLoader(file_path)
            loaded = loader.load()
            if len(loaded) < pdf_page_count:
                existing_pages = {d.metadata.get("page") for d in loaded if d.metadata.get("page") is not None}
                for page_num in range(1, pdf_page_count + 1):
                    if page_num not in existing_pages:
                        page = reader.pages[page_num - 1]
                        text = page.extract_text() or ""
                        loaded.append(Document(page_content=text, metadata={"source": file_path, "page": page_num}))
            return loaded
        finally:
            if reader is not None:
                reader.close()
    if ext in (".txt", ".md"):
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        return [Document(page_content=text, metadata={"source": file_path, "page": 1})]
    if ext == ".docx":
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            parts = []
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    parts.append(para.text)
            text = "\n\n".join(parts)
            if not text.strip():
                for table in doc.tables:
                    for row in table.rows:
                        cells = [c.text for c in row.cells]
                        parts.append(" | ".join(cells))
                text = "\n\n".join(parts) or "(empty)"
            return [Document(page_content=text, metadata={"source": file_path, "page": 1})]
        except Exception as e:
            st.warning(f"Error loading {file_path}: {e}")
            return []
    return []


def load_and_process_documents(data_folder_path: str, chunk_size: int, chunk_overlap: int) -> List[Document]:
    """Load all supported documents from the knowledge-base folder (PDF, TXT, MD, DOCX)."""
    documents = []
    if not os.path.exists(data_folder_path):
        raise FileNotFoundError(f"Knowledge base folder not found: {data_folder_path}")
    files = [f for f in os.listdir(data_folder_path) if f.lower().endswith(SUPPORTED_EXTENSIONS)]
    if not files:
        raise ValueError(
            f"No supported documents found in {data_folder_path}. "
            f"Add at least one file: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    for file in sorted(files):
        file_path = os.path.join(data_folder_path, file)
        try:
            documents.extend(_load_one_document(file_path))
        except Exception as e:
            st.warning(f"Error loading {file}: {e}")

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n\n", "\n\n", "\n", ". ", " ", ""],
    )
    return text_splitter.split_documents(documents)


def initialize_vectorstore(splits: List[Document], embeddings_model) -> FAISS:
    return FAISS.from_documents(documents=splits, embedding=embeddings_model)


def _embedding_device() -> str:
    """Use GPU for embeddings if available and not forced to CPU."""
    if EMBEDDING_DEVICE == "cpu":
        return "cpu"
    if EMBEDDING_DEVICE == "cuda":
        return "cuda"
    try:
        import torch
        return "cuda" if torch.cuda.is_available() else "cpu"
    except ImportError:
        return "cpu"


@st.cache_resource(show_spinner="Loading embeddings...")
def get_embeddings():
    device = _embedding_device()
    return HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL,
        model_kwargs={"device": device},
        encode_kwargs={"normalize_embeddings": True},
    )


@st.cache_resource(show_spinner="Loading vector store...")
def get_vectorstore(fast_mode: bool):
    """Build or load vector store. fast_mode=True uses FAST chunk settings and .vectorstore_fast cache."""
    embeddings = get_embeddings()
    if fast_mode:
        cache_path = VECTORSTORE_CACHE_PATH_FAST
        chunk_size, chunk_overlap = CHUNK_SIZE_FAST, CHUNK_OVERLAP_FAST
    else:
        cache_path = VECTORSTORE_CACHE_PATH
        chunk_size, chunk_overlap = CHUNK_SIZE_COMPLETE, CHUNK_OVERLAP_COMPLETE
    cached = _try_load_cached_vectorstore(embeddings, cache_path, chunk_size, chunk_overlap)
    if cached is not None:
        return cached
    splits = load_and_process_documents(DATA_FOLDER_PATH, chunk_size, chunk_overlap)
    vectorstore = initialize_vectorstore(splits, embeddings)
    _save_vectorstore_cache(vectorstore, cache_path, chunk_size, chunk_overlap)
    return vectorstore


@st.cache_resource(show_spinner="Loading LLM...")
def get_llm(use_ollama: bool, fast_mode: bool, groq_key: str, groq_model: str, groq_max_tokens: int,
            ollama_url: str, ollama_model: str, ollama_num_predict: int, temperature: float):
    """Return Groq or Ollama LLM; cache keyed by (use_ollama, fast_mode, models)."""
    if use_ollama:
        return ChatOllama(
            base_url=ollama_url,
            model=ollama_model,
            temperature=temperature,
            num_predict=ollama_num_predict,
        )
    return ChatGroq(
        model=groq_model,
        temperature=temperature,
        groq_api_key=groq_key,
        max_tokens=groq_max_tokens,
    )


def get_chain(use_ollama: bool, fast_mode: bool):
    """Build or return chain. Rebuilds when use_ollama or fast_mode changes."""
    key = f"chain_{'ollama' if use_ollama else 'groq'}_{'fast' if fast_mode else 'complete'}"
    if key in st.session_state:
        return st.session_state[key]

    vectorstore = get_vectorstore(fast_mode)
    if fast_mode:
        retriever = vectorstore.as_retriever(
            search_type="similarity",
            search_kwargs={"k": RETRIEVER_K_FAST},
        )
        prompt = QA_PROMPT_FAST
        groq_model, groq_max = GROQ_MODEL_FAST, GROQ_MAX_TOKENS_FAST
        ollama_model, ollama_num = OLLAMA_MODEL_FAST, OLLAMA_NUM_PREDICT_FAST
    else:
        retriever = vectorstore.as_retriever(
            search_type="mmr",
            search_kwargs={"k": RETRIEVER_K_COMPLETE, "fetch_k": RETRIEVER_FETCH_K_COMPLETE, "lambda_mult": 0.7},
        )
        prompt = QA_PROMPT
        groq_model, groq_max = GROQ_MODEL_COMPLETE, GROQ_MAX_TOKENS_COMPLETE
        ollama_model, ollama_num = OLLAMA_MODEL_COMPLETE, OLLAMA_NUM_PREDICT_COMPLETE

    memory = ConversationBufferMemory(
        memory_key="chat_history",
        output_key="answer",
        chat_memory=ChatMessageHistory(),
        return_messages=True,
    )
    llm = get_llm(use_ollama, fast_mode, GROQ_API_KEY, groq_model, groq_max,
                 OLLAMA_BASE_URL, ollama_model, ollama_num, TEMPERATURE)
    chain = ConversationalRetrievalChain.from_llm(
        llm=llm,
        chain_type="stuff",
        retriever=retriever,
        memory=memory,
        return_source_documents=True,
        verbose=False,
        combine_docs_chain_kwargs={"prompt": prompt},
    )
    st.session_state[key] = chain
    return chain


class GroqRateLimitError(Exception):
    """Raised when Groq returns 429 after all retries."""
    pass


def invoke_chain_with_timeout_and_retry(chain, question: str, timeout_sec: int, use_ollama: bool):
    """
    Run chain.invoke with timeout. For Groq (not Ollama), retry on 429 with exponential backoff (2s, 5s, 10s).
    If still 429 after retries, raise GroqRateLimitError so caller can fallback to Ollama.
    """
    def _invoke():
        return chain.invoke({"question": question})

    if use_ollama:
        with ThreadPoolExecutor(max_workers=1) as ex:
            future = ex.submit(_invoke)
            return future.result(timeout=timeout_sec)

    last_err = None
    for delay in GROQ_RETRY_DELAYS:
        try:
            with ThreadPoolExecutor(max_workers=1) as ex:
                future = ex.submit(_invoke)
                return future.result(timeout=timeout_sec)
        except Exception as e:
            last_err = e
            err_str = str(e)
            if "429" in err_str or "rate limit" in err_str.lower() or "rate_limit" in err_str.lower():
                time.sleep(delay)
                continue
            raise
    raise GroqRateLimitError(last_err)


def _run_chain_streaming(chain, question: str, timeout_sec: int, use_ollama: bool,
                         message_placeholder) -> Tuple[str, List]:
    """
    Retrieve docs, stream LLM response into message_placeholder, update memory. Returns (full_answer, source_documents).
    """
    retriever = chain.retriever
    combine_chain = chain.combine_docs_chain
    llm = combine_chain.llm
    prompt_template = combine_chain.prompt
    memory = chain.memory

    docs = retriever.invoke(question)
    context = "\n\n".join(d.page_content for d in docs)
    prompt_str = prompt_template.format(context=context, question=question)
    msg = HumanMessage(content=prompt_str)

    full_answer = ""
    try:
        for chunk in llm.stream([msg]):
            if hasattr(chunk, "content") and chunk.content:
                full_answer += chunk.content
                message_placeholder.markdown(full_answer + " ")
    except Exception:
        pass
    message_placeholder.markdown(full_answer)

    memory.chat_memory.add_user_message(question)
    memory.chat_memory.add_ai_message(full_answer)
    return full_answer, docs


# ============================================================================
# Streamlit UI
# ============================================================================

st.set_page_config(
    page_title="Knowledge Assistant",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={"Get help": None, "Report a bug": None, "About": None},
)

if "messages" not in st.session_state:
    st.session_state.messages = [
        {"role": "assistant", "content": (
            "Hello. I’m your **Knowledge Assistant**. I answer only from the documents in your knowledge base—no outside information.\n\n"
            "• **Accurate** — Every answer is grounded in your uploaded files (PDF, TXT, MD, DOCX).\n"
            "• **Traceable** — Each reply includes source references you can expand.\n"
            "• **Conversational** — You can ask follow-up questions; I keep context.\n\n"
            "Ask your question below."
        )},
    ]

# Sidebar: Mode (FAST / COMPLETE) + Model (Ollama / Groq)
with st.sidebar:
    st.markdown("### Mode")
    mode_choice = st.radio(
        "Mode",
        ["FAST (Laptop)", "COMPLETE (Server/API)"],
        index=0,
        key="mode_radio",
        help="FAST: smaller chunks, fewer docs, shorter answers. COMPLETE: full quality.",
    )
    fast_mode = mode_choice == "FAST (Laptop)"
    if "fast_mode" not in st.session_state or st.session_state.fast_mode != fast_mode:
        st.session_state.fast_mode = fast_mode
        for k in list(st.session_state.keys()):
            if k.startswith("chain_"):
                del st.session_state[k]
    st.caption("FAST = weak CPU; COMPLETE = paid API / strong server")
    st.divider()
    st.markdown("### Model")
    if "llm_choice" not in st.session_state:
        st.session_state.llm_choice = os.getenv("USE_OLLAMA", "").strip().lower() in ("1", "true", "yes")
    choice = st.radio(
        "Use",
        ["Local (Ollama)", "API (Groq)"],
        index=0 if st.session_state.llm_choice else 1,
        key="llm_radio",
    )
    use_ollama = choice == "Local (Ollama)"
    if use_ollama != st.session_state.llm_choice:
        st.session_state.llm_choice = use_ollama
        for k in list(st.session_state.keys()):
            if k.startswith("chain_"):
                del st.session_state[k]
    if use_ollama:
        model_name = OLLAMA_MODEL_FAST if fast_mode else OLLAMA_MODEL_COMPLETE
        st.caption(f"**{model_name}** (local)")
        st.caption("[Ollama](https://ollama.com)")
    else:
        model_name = GROQ_MODEL_FAST if fast_mode else GROQ_MODEL_COMPLETE
        st.caption(f"**{model_name}** (API)")
        if not GROQ_API_KEY or GROQ_API_KEY == "your-groq-api-key-here":
            st.warning("Set GROQ_API_KEY in .env to use Groq.")
    st.divider()
    st.caption("Answers include **source excerpts** below each reply.")

# Initialize RAG for current mode (load vector store)
try:
    get_vectorstore(st.session_state.get("fast_mode", False))
except Exception as e:
    st.error(f"Cannot start RAG: {e}")
    st.stop()

# Chat
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])
        if msg.get("sources"):
            with st.expander("Sources"):
                for s in msg["sources"]:
                    st.markdown(s)
                    st.divider()

if prompt := st.chat_input("Ask about your documents..."):
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    fast_mode = st.session_state.get("fast_mode", False)
    use_ollama = st.session_state.llm_choice
    timeout_sec = LLM_TIMEOUT_FAST if fast_mode else LLM_TIMEOUT_COMPLETE

    with st.chat_message("assistant"):
        source_texts = []
        message_placeholder = st.empty()
        try:
            chain = get_chain(use_ollama, fast_mode)
            # Prefer streaming so the UI doesn't look stuck
            try:
                message_placeholder.markdown("Thinking...")
                answer, source_documents = _run_chain_streaming(
                    chain, prompt, timeout_sec, use_ollama, message_placeholder
                )
            except Exception:
                # Fallback: non-streaming invoke with retry (Groq 429 -> backoff, then optional Ollama fallback)
                message_placeholder.markdown("Thinking...")
                res = invoke_chain_with_timeout_and_retry(chain, prompt, timeout_sec, use_ollama)
                answer = res["answer"]
                source_documents = res.get("source_documents") or []
                message_placeholder.markdown(answer)
        except GroqRateLimitError:
            message_placeholder.markdown(
                "**Rate limit (Groq)** after retries. Falling back to **Local (Ollama)** for this reply."
            )
            use_ollama = True
            st.session_state.llm_choice = True
            for k in list(st.session_state.keys()):
                if k.startswith("chain_"):
                    del st.session_state[k]
            chain = get_chain(use_ollama, fast_mode)
            message_placeholder.markdown("Thinking...")
            try:
                res = invoke_chain_with_timeout_and_retry(chain, prompt, timeout_sec, use_ollama)
                answer = res["answer"]
                source_documents = res.get("source_documents") or []
                message_placeholder.markdown(answer)
            except (FuturesTimeoutError, Exception) as e2:
                answer = f"**Error:** {e2}\n\nTry again or switch mode in the sidebar."
                source_documents = []
                message_placeholder.markdown(answer)
        except FuturesTimeoutError:
            answer = (
                "**Request timed out.** Try **FAST (Laptop)** mode or **Local (Ollama)** in the sidebar."
            )
            source_documents = []
            message_placeholder.markdown(answer)
        except Exception as e:
            err = str(e)
            if "429" in err or "rate limit" in err.lower():
                answer = "**Rate limit (Groq).** Switch to **Local (Ollama)** or try again later."
            else:
                answer = f"**Error:** {err}\n\nTry switching Mode or Model in the sidebar."
            source_documents = []
            message_placeholder.markdown(answer)

        if source_documents:
            for doc in source_documents:
                page = doc.metadata.get("page", "N/A")
                src_file = doc.metadata.get("source", "Unknown")
                preview = doc.page_content
                if len(preview) > MAX_SOURCE_PREVIEW:
                    preview = preview[:MAX_SOURCE_PREVIEW].rstrip() + "..."
                source_texts.append(f"**Page {page}** | {os.path.basename(src_file)}\n\n{preview}")
            answer += "\n\n---\n*Sources: expand section below.*"

        if source_texts:
            with st.expander("Sources"):
                for s in source_texts:
                    st.markdown(s)
                    st.divider()

    st.session_state.messages.append({"role": "assistant", "content": answer, "sources": source_texts})
